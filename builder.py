import os
import docx

FILE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_DIR = os.path.join(FILE_DIR, "Template.docx")
DEST_DIR = os.path.join(FILE_DIR, "file.docx")
SNIPPETS_DIR = os.path.join(FILE_DIR, "Snippets.docx")

# os.system(f"cp {TEMPLATE_DIR} {DEST_DIR}")

doc_src = docx.Document(SNIPPETS_DIR)
doc_dest = docx.Document(DEST_DIR)
skills = []
paragraphs = []

index = 0
for para in doc_src.paragraphs:
    if para.text != "":
        paragraphs.append(para.text)
    for run in para.runs:
        if run.bold:
            print(f"[{index}]: {run.text}")
            skills.append(run.text)
            index += 1

desired_skills = input("Enter a comma seperated list of desired skills: ")
tags = desired_skills.split(",")
target_para = []

for s in tags:
    desired_index = int(s)
    target_para.append(paragraphs[desired_index * 2 + 1])

text_add = "\n\n".join(target_para)

for para in doc_dest.paragraphs:
    if "</Paragraphs>" in para:
        para.text = text_add
        break
